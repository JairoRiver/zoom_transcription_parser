from docx import Document

def parse_and_save(sections, save_name):

    # Crea un nuevo documento
    doc = Document()

    for section in sections:
        text = section['transcription']
        indice_usuario = text.find(":")
        parte1 = text[:indice_usuario + 1]
        parte2 = text[indice_usuario + 1:]

        # Añade un párrafo con la primera parte en negrita
        paragraph = doc.add_paragraph()
        run1 = paragraph.add_run(parte1)
        run1.bold = True

        # Añade el resto del texto (parte2)
        paragraph.add_run(parte2)

        # Agrega una línea en blanco al final del parrafo
        doc.add_paragraph()

    # Guarda el documento en un archivo
    
    doc.save(save_name)
